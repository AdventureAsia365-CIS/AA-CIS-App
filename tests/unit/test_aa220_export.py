"""AA-220: multi-version export (A), single-version DOCX (B), multi-tour filter (D).

Pure tests — no DB / no AWS. The real endpoint functions are called with a fake
pool/conn (same pattern as test_aa219_version_detail), so the transpose builder,
the DOCX section assembler, and the tour_ids WHERE-param construction are all
exercised as shipped.
"""

import io
import json
import pytest
from unittest.mock import patch

from docx import Document

from api.routers import admin_pipeline


# ── fake pool / conn / request ────────────────────────────────────────────────
class _FakeConn:
    def __init__(self, *, rows=None, row=None, capture=None):
        self._rows = rows
        self._row = row
        self._capture = capture  # list to record (query, args)

    async def fetch(self, query, *args):
        if self._capture is not None:
            self._capture.append((query, args))
        return self._rows if self._rows is not None else []

    async def fetchrow(self, query, *args):
        return self._row


class _FakePool:
    def __init__(self, conn):
        self._conn = conn

    def acquire(self):
        conn = self._conn

        class _Ctx:
            async def __aenter__(self):
                return conn

            async def __aexit__(self, *a):
                return False
        return _Ctx()


class _FakeRequest:
    def __init__(self, pool):
        self.app = type("A", (), {"state": type("S", (), {"pool": pool})()})()


def _request(conn):
    return _FakeRequest(_FakePool(conn))


async def _read_body(resp):
    out = b""
    async for chunk in resp.body_iterator:
        out += chunk if isinstance(chunk, (bytes, bytearray)) else chunk.encode("utf-8")
    return out


_NOSECRET = patch.object(admin_pipeline, "verify_admin_secret", lambda *_a, **_k: None)
TOUR = "ca893afe-27e2-431b-9596-b92514e7f98c"


# ── STAGE A — multi-version transpose (field × version) ───────────────────────
def _a_row(version_num, model_id, judge_score):
    return {
        "version_num": version_num, "model_id": model_id,
        "score_overall": 10.0, "score_brand": 9.0, "score_seo": 8.5,
        "score_structure": 10.0, "score_quality": 9.5,
        "failure_codes": json.dumps(["SEO_META_SHORT"]),
        "brand_audit_status": "pass",
        "brand_audit_codes": json.dumps([]), "brand_audit_issues": [],
        "fix_pass_applied": True, "fix_pass_fields": json.dumps(["seo_meta"]),
        "metadata": json.dumps({"judge": {"brand_fit": 8.0, "distinct": 7.0,
                                          "mission_present": True, "feedback": "fb",
                                          "judge_score": judge_score},
                                "llm_cost_usd": 0.002}),
        "aa_name": "Seoul Executive", "aa_subtitle": "sub", "aa_summary": "summ",
        "aa_description": "desc", "aa_highlights": json.dumps(["h1", "h2"]),
        "aa_itineraries": "Day 1\nDay 2", "seo_title": "t", "seo_meta": "m",
        "brand_name": "AA",
    }


@pytest.mark.asyncio
async def test_a_transpose_field_by_version():
    rows = [_a_row(6, "haiku", 7.0), _a_row(7, "sonnet", 9.0)]
    with _NOSECRET:
        resp = await admin_pipeline.export_tour_versions(
            TOUR, _request(_FakeConn(rows=rows)), versions="6,7", format="csv",
            x_admin_secret="s",
        )
    text = (await _read_body(resp)).decode("utf-8")
    lines = [ln.strip("\r") for ln in text.strip().split("\n")]

    # header row: field label + one column per version
    assert lines[0] == "field,v6,v7"
    # each field is a ROW, each version a COLUMN (true transpose)
    assert "version_num,6,7" in lines
    assert "model,haiku,sonnet" in lines
    assert "brand_name,AA,AA" in lines
    assert "judge_score,7.0,9.0" in lines       # pulled from metadata.judge per version
    assert "score_overall,10.0,10.0" in lines


@pytest.mark.asyncio
async def test_a_xlsx_returns_spreadsheet_bytes():
    rows = [_a_row(6, "haiku", 7.0)]
    with _NOSECRET:
        resp = await admin_pipeline.export_tour_versions(
            TOUR, _request(_FakeConn(rows=rows)), versions="6", format="xlsx",
            x_admin_secret="s",
        )
    body = await _read_body(resp)
    assert body[:2] == b"PK"  # xlsx is a zip
    assert "spreadsheetml" in resp.media_type


@pytest.mark.asyncio
async def test_a_bad_inputs_raise_400_and_missing_raises_404():
    with _NOSECRET:
        with pytest.raises(admin_pipeline.HTTPException) as e1:
            await admin_pipeline.export_tour_versions(
                TOUR, _request(_FakeConn(rows=[])), versions="", format="csv", x_admin_secret="s")
        assert e1.value.status_code == 400

        with pytest.raises(admin_pipeline.HTTPException) as e2:
            await admin_pipeline.export_tour_versions(
                TOUR, _request(_FakeConn(rows=[])), versions="abc", format="csv", x_admin_secret="s")
        assert e2.value.status_code == 400

        with pytest.raises(admin_pipeline.HTTPException) as e3:
            await admin_pipeline.export_tour_versions(
                TOUR, _request(_FakeConn(rows=[])), versions="9", format="csv", x_admin_secret="s")
        assert e3.value.status_code == 404


# ── STAGE B — DOCX section assembler ──────────────────────────────────────────
def _b_row():
    return {
        "id": "e46d69a9-6721-4aac-8732-59d863374b94",
        "tenant_id": "00000000-0000-0000-0000-000000000001",
        "version_num": 6, "model_id": "haiku",
        "score_overall": 10.0, "score_brand": 9.0, "score_seo": 8.5,
        "score_structure": 10.0, "score_quality": 9.5,
        "failure_codes": json.dumps(["SEO_META_SHORT"]),
        "brand_audit_status": "flagged",
        "brand_audit_codes": json.dumps(["FORBIDDEN_WORD"]),
        "brand_audit_issues": [{"field": "seo_meta", "msg": "x"}],
        "fix_pass_applied": True, "fix_pass_fields": json.dumps(["seo_meta"]),
        "created_at": None,
        "aa_name": "Seoul Executive Discovery", "aa_subtitle": "sub",
        "aa_summary": "summary text", "aa_description": "d",
        "aa_highlights": json.dumps(["Private guide", "Helicopter transfer"]),
        "aa_itineraries": "Day 1: Arrival\nDay 2: Temple\nDay 3: Depart",
        "seo_title": "Executive Korea", "seo_meta": "A discreet luxury journey.",
        "metadata": json.dumps({"judge": {"brand_fit": 8.0, "distinct": 7.0,
                                          "mission_present": True, "feedback": "Lead w/ privacy.",
                                          "judge_score": 7.0},
                                "revalidate_ran": True, "revalidate_passed": True}),
        "brand_name": "AA",
        "top_keywords": json.dumps(["korea tours"]),
        "keyword_ideas": [{"keyword": "korea trek", "search_volume": 320,
                           "competition": "LOW", "competition_index": 5, "cpc": 0.93}],
        "people_also_ask": json.dumps(["is korea safe?"]),
        "related_keywords": ["korea tours", "seoul trip"],
        "keyword_search": "South Korea tours",
    }


async def _build_docx(row):
    with _NOSECRET:
        resp = await admin_pipeline.export_tour_version_docx(
            TOUR, 6, _request(_FakeConn(row=row)), x_admin_secret="s")
    return resp, Document(io.BytesIO(await _read_body(resp)))


@pytest.mark.asyncio
async def test_b_docx_has_all_sections():
    resp, doc = await _build_docx(_b_row())
    assert "wordprocessingml.document" in resp.media_type
    texts = [p.text for p in doc.paragraphs]
    joined = "\n".join(texts)
    for section in ("Scores & Judge", "Brand Audit", "SEO / DataForSEO",
                    "Content", "Summary", "Highlights", "Itineraries", "SEO"):
        assert any(section in t for t in texts), f"missing section: {section}"
    # header + brand + judge feedback surfaced
    assert "Seoul Executive Discovery" in joined
    assert "Lead w/ privacy." in joined


@pytest.mark.asyncio
async def test_b_dfs_table_and_seo_meta_charcount():
    _resp, doc = await _build_docx(_b_row())
    # DFS keyword_ideas table: header + 1 data row, 5 cols
    assert len(doc.tables) == 1
    tbl = doc.tables[0]
    header = [c.text for c in tbl.rows[0].cells]
    assert header == ["Keyword", "Volume", "Competition", "Comp. Index", "CPC"]
    data = [c.text for c in tbl.rows[1].cells]
    assert data[0] == "korea trek" and data[1] == "320"
    # seo_meta char count row
    joined = "\n".join(p.text for p in doc.paragraphs)
    assert f"{len('A discreet luxury journey.')} chars" in joined


@pytest.mark.asyncio
async def test_b_itineraries_keep_newlines_not_flattened():
    _resp, doc = await _build_docx(_b_row())
    joined = "\n".join(p.text for p in doc.paragraphs)
    # both itinerary lines present (substrings survive even though runs concat in .text)
    assert "Day 1: Arrival" in joined
    assert "Day 3: Depart" in joined
    # the soft line breaks were actually emitted as <w:br/> (not flattened to one run)
    xml = "".join(p._p.xml for p in doc.paragraphs)
    assert "w:br" in xml


@pytest.mark.asyncio
async def test_b_missing_version_raises_404():
    with _NOSECRET:
        with pytest.raises(admin_pipeline.HTTPException) as e:
            await admin_pipeline.export_tour_version_docx(
                TOUR, 99, _request(_FakeConn(row=None)), x_admin_secret="s")
        assert e.value.status_code == 404


# ── STAGE D — tour_ids parse + WHERE-param construction ───────────────────────
async def _capture_export_tours(tour_ids):
    cap = []
    with _NOSECRET:
        await admin_pipeline.export_tours(
            _request(_FakeConn(rows=[], capture=cap)),
            format="csv", tour_ids=tour_ids, x_admin_secret="s")
    return cap[0]  # (query, args)


@pytest.mark.asyncio
async def test_d_no_tour_ids_means_no_where_and_no_params():
    query, args = await _capture_export_tours(None)
    assert "WHERE" not in query
    assert args == ()


@pytest.mark.asyncio
async def test_d_tour_ids_builds_where_any_and_parses_csv():
    query, args = await _capture_export_tours(" id1, id2 ,,id3 ")
    assert "WHERE rt.tour_id = ANY($1::uuid[])" in query
    # parsed: trimmed, empties dropped, order preserved
    assert args == (["id1", "id2", "id3"],)
    # filter sits BEFORE the GROUP BY (pre-aggregation)
    assert query.index("ANY($1::uuid[])") < query.index("GROUP BY")


@pytest.mark.asyncio
async def test_d_blank_tour_ids_degrades_to_no_filter():
    query, args = await _capture_export_tours("  , ,")
    assert "WHERE" not in query
    assert args == ()
